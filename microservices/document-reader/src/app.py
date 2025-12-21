"""
Document Reader Microservice
Processes documents, extracts text, generates notes using AI, and publishes to Kafka
"""
import os
import json
import logging
import uuid
import signal
import sys
from datetime import datetime
from flask import Flask, request, jsonify
from flask_cors import CORS
import boto3
from botocore.exceptions import ClientError
import psycopg2
from psycopg2.extras import RealDictCursor
from kafka import KafkaProducer
from kafka.errors import KafkaError
import time
import PyPDF2
import io
from docx import Document
import requests

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

app = Flask(__name__)
CORS(app)

# Environment configuration
AWS_REGION = os.environ.get('AWS_REGION', 'us-east-1')
S3_BUCKET = os.environ.get('S3_BUCKET_DOCUMENTS', 'document-reader-storage-dev-653040176723')
DB_HOST = os.environ.get('DB_HOST', 'cloud-project-db.cjyjymrymgig.us-east-1.rds.amazonaws.com')
DB_PORT = os.environ.get('DB_PORT', '5432')
DB_NAME = os.environ.get('DB_NAME', 'document_reader')
DB_USER = os.environ.get('DB_USER', 'postgres')
DB_PASSWORD = os.environ.get('DB_PASSWORD', 'MySecurePassword123!')
KAFKA_BOOTSTRAP_SERVERS = os.environ.get('KAFKA_BOOTSTRAP_SERVERS', 'localhost:9092').split(',')
KAFKA_TOPIC_UPLOADED = 'document.uploaded'
KAFKA_TOPIC_PROCESSED = 'document.processed'
KAFKA_TOPIC_NOTES = 'notes.generated'

# OpenRouter configuration
OPENROUTER_API_KEY = os.environ.get('OPENROUTER_API_KEY', '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')

# Initialize AWS clients
try:
    s3_client = boto3.client('s3', region_name=AWS_REGION)
    textract_client = boto3.client('textract', region_name=AWS_REGION)
    logger.info("AWS clients initialized successfully")
except Exception as e:
    logger.error(f"Failed to initialize AWS clients: {e}")
    s3_client = None
    textract_client = None

# Initialize Kafka producer
def create_kafka_producer():
    """Create Kafka producer with retry logic"""
    max_retries = 5
    retry_delay = 2
    
    for attempt in range(max_retries):
        try:
            producer = KafkaProducer(
                bootstrap_servers=KAFKA_BOOTSTRAP_SERVERS,
                value_serializer=lambda v: json.dumps(v).encode('utf-8'),
                acks='all',
                retries=3,
                max_in_flight_requests_per_connection=1
            )
            logger.info("Kafka producer initialized successfully")
            return producer
        except KafkaError as e:
            logger.warning(f"Kafka connection attempt {attempt + 1} failed: {e}")
            if attempt < max_retries - 1:
                time.sleep(retry_delay)
            else:
                logger.error("Failed to initialize Kafka producer after all retries")
                return None

kafka_producer = create_kafka_producer()

# Initialize OpenRouter API
openrouter_available = bool(OPENROUTER_API_KEY)
if openrouter_available:
    logger.info("OpenRouter API configured successfully")
else:
    logger.warning("OpenRouter API key not configured")


def signal_handler(sig, frame):
    """Handle shutdown signals for graceful termination"""
    logger.info(f"Received signal {sig}, initiating graceful shutdown...")
    
    # Close Kafka producer
    if kafka_producer:
        kafka_producer.close()
        logger.info("Kafka producer closed")
    
    sys.exit(0)


# Register signal handlers for graceful shutdown
signal.signal(signal.SIGTERM, signal_handler)
signal.signal(signal.SIGINT, signal_handler)


def get_db_connection():
    """Get PostgreSQL database connection"""
    try:
        conn = psycopg2.connect(
            host=DB_HOST,
            port=DB_PORT,
            database=DB_NAME,
            user=DB_USER,
            password=DB_PASSWORD,
            cursor_factory=RealDictCursor
        )
        return conn
    except Exception as e:
        logger.error(f"Database connection error: {e}")
        return None


# Database initialization removed - using existing AWS RDS database
# Tables are pre-configured in AWS RDS


def ensure_s3_bucket():
    """Ensure S3 bucket exists"""
    if not s3_client:
        return False
    
    try:
        s3_client.head_bucket(Bucket=S3_BUCKET)
        logger.info(f"S3 bucket {S3_BUCKET} exists")
        return True
    except ClientError as e:
        error_code = e.response['Error']['Code']
        if error_code == '404':
            try:
                s3_client.create_bucket(Bucket=S3_BUCKET)
                logger.info(f"Created S3 bucket: {S3_BUCKET}")
                return True
            except ClientError as ce:
                logger.error(f"Failed to create S3 bucket: {ce}")
                return False
        else:
            logger.error(f"Error checking S3 bucket: {e}")
            return False


def extract_text_from_pdf(file_data):
    """
    Extract text from PDF file
    
    Args:
        file_data: PDF file binary data
    
    Returns:
        str: Extracted text or None on error
    """
    try:
        pdf_file = io.BytesIO(file_data)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        text_content = []
        for page in pdf_reader.pages:
            text_content.append(page.extract_text())
        
        extracted_text = '\n'.join(text_content)
        logger.info(f"Extracted {len(extracted_text)} characters from PDF")
        
        return extracted_text
        
    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        return None


def extract_text_from_docx(file_data):
    """
    Extract text from DOCX file
    
    Args:
        file_data: DOCX file binary data
    
    Returns:
        str: Extracted text or None on error
    """
    try:
        docx_file = io.BytesIO(file_data)
        document = Document(docx_file)
        
        text_content = []
        # Extract text from paragraphs
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Extract text from tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_content.append(cell.text)
        
        extracted_text = '\n'.join(text_content)
        logger.info(f"Extracted {len(extracted_text)} characters from DOCX")
        
        return extracted_text
        
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        return None


def generate_notes_from_text(text, max_length=1000):
    """
    Generate notes from text using OpenRouter AI
    
    Args:
        text: Document text
        max_length: Maximum note length
    
    Returns:
        str: Generated notes
    """
    # Use OpenRouter API if available
    if openrouter_available:
        try:
            # Truncate text if too long
            max_input_chars = 12000  # Roughly 3000 tokens
            input_text = text[:max_input_chars] if len(text) > max_input_chars else text
            
            prompt = f"""Analyze the following document and create comprehensive study notes. 
Include:
1. Main topics and key concepts
2. Important details and definitions
3. Key takeaways
4. Summary of main points

Document:
{input_text}

Provide well-structured notes:"""
            
            headers = {
                "Authorization": f"Bearer {OPENROUTER_API_KEY}",
                "Content-Type": "application/json"
            }
            
            payload = {
                "model": "amazon/nova-2-lite-v1:free",
                "messages": [
                    {"role": "system", "content": "You are an expert at creating clear, concise study notes from documents. Create organized, easy-to-understand notes."},
                    {"role": "user", "content": prompt}
                ],
                "max_tokens": 1500,
                "temperature": 0.3
            }
            
            response = requests.post(
                "https://openrouter.ai/api/v1/chat/completions",
                headers=headers,
                json=payload,
                timeout=30
            )
            
            if response.status_code == 200:
                result = response.json()
                notes = result['choices'][0]['message']['content']
                logger.info(f"Generated AI notes: {len(notes)} characters")
                return notes
            else:
                logger.error(f"OpenRouter API error: {response.status_code} - {response.text}")
                # Fall back to simple summary
            
        except Exception as e:
            logger.error(f"AI note generation error: {e}")
            # Fall back to simple summary
    
    # Fallback: Simple summary if AI is not available
    if len(text) <= max_length:
        notes = f"Document Summary:\n\n{text}"
    else:
        # Create a basic summary
        lines = text.split('\n')
        summary_lines = []
        char_count = 0
        
        for line in lines:
            if char_count + len(line) > max_length:
                break
            if line.strip():
                summary_lines.append(line)
                char_count += len(line)
        
        notes = f"Document Summary:\n\n" + '\n'.join(summary_lines)
        if len(text) > max_length:
            notes += "\n\n[Content truncated for brevity]"
    
    return notes


def publish_to_kafka(topic, message):
    """
    Publish message to Kafka topic
    
    Args:
        topic: Kafka topic name
        message: Message dictionary
    
    Returns:
        bool: True if successful, False otherwise
    """
    if not kafka_producer:
        logger.error("Kafka producer not initialized")
        return False
    
    try:
        future = kafka_producer.send(topic, value=message)
        record_metadata = future.get(timeout=10)
        
        logger.info(f"Message published to {topic}: partition={record_metadata.partition}, offset={record_metadata.offset}")
        return True
        
    except KafkaError as e:
        logger.error(f"Kafka publish error: {e}")
        return False
    except Exception as e:
        logger.error(f"Unexpected error publishing to Kafka: {e}")
        return False


@app.route('/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    db_conn = get_db_connection()
    db_healthy = db_conn is not None
    if db_conn:
        db_conn.close()
    
    health_status = {
        'service': 'document-reader',
        'status': 'healthy' if db_healthy and s3_client and kafka_producer else 'unhealthy',
        'timestamp': datetime.utcnow().isoformat(),
        'components': {
            's3': s3_client is not None,
            'textract': textract_client is not None,
            'database': db_healthy,
            'kafka': kafka_producer is not None,
            'openrouter': openrouter_available
        }
    }
    
    status_code = 200 if all(health_status['components'].values()) else 503
    return jsonify(health_status), status_code


@app.route('/ready', methods=['GET'])
def readiness_check():
    """Readiness check endpoint"""
    return jsonify({'status': 'ready'}), 200


@app.route('/api/documents/upload', methods=['POST'])
def upload_document():
    """
    Upload document
    
    Request: multipart/form-data
    - file: Document file (required)
    - user_id: User ID (required)
    """
    try:
        # Validate request
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        user_id = request.form.get('user_id')
        
        if not user_id:
            return jsonify({'error': 'Missing required field: user_id'}), 400
        
        if file.filename == '':
            return jsonify({'error': 'Empty filename'}), 400
        
        # Validate file type
        allowed_extensions = {'pdf', 'txt', 'doc', 'docx'}
        file_extension = file.filename.rsplit('.', 1)[1].lower() if '.' in file.filename else ''
        
        if file_extension not in allowed_extensions:
            return jsonify({'error': f'File type not supported. Allowed: {allowed_extensions}'}), 400
        
        # Read file data
        file_data = file.read()
        file_size = len(file_data)
        
        # Validate file size (max 50MB)
        if file_size > 50 * 1024 * 1024:
            return jsonify({'error': 'File too large (max 50MB)'}), 400
        
        # Generate document ID
        document_id = str(uuid.uuid4())
        
        # Ensure S3 bucket exists
        ensure_s3_bucket()
        
        # Upload to S3
        file_key = f"documents/{user_id}/{document_id}/{file.filename}"
        
        try:
            s3_client.put_object(
                Bucket=S3_BUCKET,
                Key=file_key,
                Body=file_data,
                ContentType=file.content_type or 'application/octet-stream'
            )
            
            s3_url = f"s3://{S3_BUCKET}/{file_key}"
            logger.info(f"Document uploaded to S3: {s3_url}")
            
        except ClientError as e:
            logger.error(f"S3 upload error: {e}")
            return jsonify({'error': 'Failed to upload document'}), 500
        
        # Extract text content
        text_content = None
        if file_extension == 'pdf':
            text_content = extract_text_from_pdf(file_data)
        elif file_extension in ['docx', 'doc']:
            text_content = extract_text_from_docx(file_data)
        elif file_extension == 'txt':
            text_content = file_data.decode('utf-8', errors='ignore')
        
        # Store in database
        conn = get_db_connection()
        if not conn:
            return jsonify({'error': 'Database connection failed'}), 503
        
        try:
            cursor = conn.cursor()
            cursor.execute(
                """
                INSERT INTO documents 
                (document_id, user_id, filename, file_size, content_type, s3_url, text_content, status)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                """,
                (document_id, user_id, file.filename, file_size, file.content_type, 
                 s3_url, text_content, 'uploaded')
            )
            conn.commit()
            cursor.close()
            conn.close()
            
        except Exception as e:
            logger.error(f"Database error: {e}")
            if conn:
                conn.close()
            return jsonify({'error': 'Failed to store document metadata'}), 500
        
        # Publish uploaded event to Kafka
        upload_event = {
            'document_id': document_id,
            'user_id': user_id,
            'filename': file.filename,
            'file_size': file_size,
            's3_url': s3_url,
            'status': 'uploaded',
            'timestamp': datetime.utcnow().isoformat()
        }
        
        publish_to_kafka(KAFKA_TOPIC_UPLOADED, upload_event)
        
        logger.info(f"Document uploaded: {document_id}")
        
        return jsonify({
            'document_id': document_id,
            'filename': file.filename,
            'file_size': file_size,
            's3_url': s3_url,
            'status': 'uploaded'
        }), 201
        
    except Exception as e:
        logger.error(f"Unexpected error in upload_document: {e}")
        return jsonify({'error': 'Internal server error'}), 500


@app.route('/api/documents/<document_id>', methods=['GET'])
def get_document(document_id):
    """Get document details"""
    try:
        conn = get_db_connection()
        if not conn:
            return jsonify({'error': 'Database connection failed'}), 503
        
        try:
            cursor = conn.cursor()
            cursor.execute(
                "SELECT * FROM documents WHERE document_id = %s",
                (document_id,)
            )
            document = cursor.fetchone()
            
            cursor.close()
            conn.close()
            
            if not document:
                return jsonify({'error': 'Document not found'}), 404
            
            return jsonify(dict(document)), 200
            
        except Exception as e:
            logger.error(f"Database error: {e}")
            if conn:
                conn.close()
            return jsonify({'error': 'Failed to retrieve document'}), 500
        
    except Exception as e:
        logger.error(f"Unexpected error in get_document: {e}")
        return jsonify({'error': 'Internal server error'}), 500


@app.route('/api/documents/<document_id>/notes', methods=['GET'])
def get_document_notes(document_id):
    """Get notes for a document"""
    try:
        conn = get_db_connection()
        if not conn:
            return jsonify({'error': 'Database connection failed'}), 503
        
        try:
            cursor = conn.cursor()
            cursor.execute(
                "SELECT * FROM notes WHERE document_id = %s ORDER BY created_at DESC",
                (document_id,)
            )
            notes = cursor.fetchall()
            
            cursor.close()
            conn.close()
            
            return jsonify({
                'document_id': document_id,
                'notes': [dict(note) for note in notes]
            }), 200
            
        except Exception as e:
            logger.error(f"Database error: {e}")
            if conn:
                conn.close()
            return jsonify({'error': 'Failed to retrieve notes'}), 500
        
    except Exception as e:
        logger.error(f"Unexpected error in get_document_notes: {e}")
        return jsonify({'error': 'Internal server error'}), 500


@app.route('/api/documents/<document_id>/regenerate-notes', methods=['POST'])
def regenerate_notes(document_id):
    """Regenerate notes for a document"""
    try:
        # Get document from database
        conn = get_db_connection()
        if not conn:
            return jsonify({'error': 'Database connection failed'}), 503
        
        try:
            cursor = conn.cursor()
            cursor.execute(
                "SELECT * FROM documents WHERE document_id = %s",
                (document_id,)
            )
            document = cursor.fetchone()
            
            if not document:
                cursor.close()
                conn.close()
                return jsonify({'error': 'Document not found'}), 404
            
            text_content = document['text_content']
            user_id = document['user_id']
            
            if not text_content:
                cursor.close()
                conn.close()
                return jsonify({'error': 'No text content available for processing'}), 400
            
            # Generate new notes
            notes_content = generate_notes_from_text(text_content)
            note_id = str(uuid.uuid4())
            
            # Store notes
            cursor.execute(
                """
                INSERT INTO notes (note_id, document_id, user_id, note_content)
                VALUES (%s, %s, %s, %s)
                """,
                (note_id, document_id, user_id, notes_content)
            )
            
            # Update document status
            cursor.execute(
                "UPDATE documents SET status = %s, updated_at = CURRENT_TIMESTAMP WHERE document_id = %s",
                ('processed', document_id)
            )
            
            conn.commit()
            cursor.close()
            conn.close()
            
            # Publish processed event
            processed_event = {
                'document_id': document_id,
                'user_id': user_id,
                'status': 'processed',
                'timestamp': datetime.utcnow().isoformat()
            }
            
            publish_to_kafka(KAFKA_TOPIC_PROCESSED, processed_event)
            
            # Publish notes generated event
            notes_event = {
                'note_id': note_id,
                'document_id': document_id,
                'user_id': user_id,
                'timestamp': datetime.utcnow().isoformat()
            }
            
            publish_to_kafka(KAFKA_TOPIC_NOTES, notes_event)
            
            logger.info(f"Notes regenerated for document: {document_id}")
            
            return jsonify({
                'document_id': document_id,
                'note_id': note_id,
                'status': 'processed'
            }), 200
            
        except Exception as e:
            logger.error(f"Database error: {e}")
            if conn:
                conn.close()
            return jsonify({'error': 'Failed to regenerate notes'}), 500
        
    except Exception as e:
        logger.error(f"Unexpected error in regenerate_notes: {e}")
        return jsonify({'error': 'Internal server error'}), 500


@app.route('/api/documents', methods=['GET'])
def list_documents():
    """List user documents"""
    try:
        user_id = request.args.get('user_id')
        
        if not user_id:
            return jsonify({'error': 'Missing required parameter: user_id'}), 400
        
        limit = request.args.get('limit', 20, type=int)
        offset = request.args.get('offset', 0, type=int)
        
        if limit > 100:
            limit = 100
        
        conn = get_db_connection()
        if not conn:
            return jsonify({'error': 'Database connection failed'}), 503
        
        try:
            cursor = conn.cursor()
            cursor.execute(
                """
                SELECT document_id, filename, file_size, status, created_at
                FROM documents
                WHERE user_id = %s
                ORDER BY created_at DESC
                LIMIT %s OFFSET %s
                """,
                (user_id, limit, offset)
            )
            documents = cursor.fetchall()
            
            cursor.execute(
                "SELECT COUNT(*) as total FROM documents WHERE user_id = %s",
                (user_id,)
            )
            total = cursor.fetchone()['total']
            
            cursor.close()
            conn.close()
            
            return jsonify({
                'documents': [dict(doc) for doc in documents],
                'total': total,
                'limit': limit,
                'offset': offset
            }), 200
            
        except Exception as e:
            logger.error(f"Database error: {e}")
            if conn:
                conn.close()
            return jsonify({'error': 'Failed to list documents'}), 500
        
    except Exception as e:
        logger.error(f"Unexpected error in list_documents: {e}")
        return jsonify({'error': 'Internal server error'}), 500


@app.route('/api/documents/<document_id>', methods=['DELETE'])
def delete_document(document_id):
    """Delete a document and its associated data"""
    try:
        conn = get_db_connection()
        if not conn:
            return jsonify({'error': 'Database connection failed'}), 503
        
        try:
            cursor = conn.cursor()
            
            # Get document details for S3 deletion
            cursor.execute(
                "SELECT s3_url, user_id FROM documents WHERE document_id = %s",
                (document_id,)
            )
            document = cursor.fetchone()
            
            if not document:
                cursor.close()
                conn.close()
                return jsonify({'error': 'Document not found'}), 404
            
            s3_url = document['s3_url']
            user_id = document['user_id']
            
            # Delete from S3
            if s3_client and s3_url.startswith('s3://'):
                try:
                    # Extract bucket and key from S3 URL
                    s3_path = s3_url.replace('s3://', '')
                    bucket, key = s3_path.split('/', 1)
                    
                    s3_client.delete_object(Bucket=bucket, Key=key)
                    logger.info(f"Deleted S3 object: {s3_url}")
                except Exception as e:
                    logger.error(f"Failed to delete S3 object: {e}")
                    # Continue with database deletion even if S3 fails
            
            # Delete notes (cascade should handle this, but explicit delete is safer)
            cursor.execute(
                "DELETE FROM notes WHERE document_id = %s",
                (document_id,)
            )
            
            # Delete document
            cursor.execute(
                "DELETE FROM documents WHERE document_id = %s",
                (document_id,)
            )
            
            conn.commit()
            cursor.close()
            conn.close()
            
            logger.info(f"Document deleted: {document_id}")
            
            return jsonify({
                'message': 'Document deleted successfully',
                'document_id': document_id
            }), 200
            
        except Exception as e:
            logger.error(f"Database error: {e}")
            if conn:
                conn.close()
            return jsonify({'error': 'Failed to delete document'}), 500
        
    except Exception as e:
        logger.error(f"Unexpected error in delete_document: {e}")
        return jsonify({'error': 'Internal server error'}), 500


@app.errorhandler(404)
def not_found(error):
    """Handle 404 errors"""
    return jsonify({'error': 'Endpoint not found'}), 404


@app.errorhandler(500)
def internal_error(error):
    """Handle 500 errors"""
    return jsonify({'error': 'Internal server error'}), 500


if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5004))
    app.run(host='0.0.0.0', port=port, debug=False)
